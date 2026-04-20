from __future__ import annotations

import json
import math
from collections import Counter
from pathlib import Path
from typing import Any

from django.conf import settings
from django.core.management.base import BaseCommand

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

from app.assignment import AssignmentPipelineService


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _table_style(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"


def _safe_text(value: Any) -> str:
    if isinstance(value, float):
        return f"{value:.4f}"
    return str(value)


class Command(BaseCommand):
    help = "Xuat file Word cho bai nop AISERVICE."

    def handle(self, *args, **options):
        service = AssignmentPipelineService()
        summary = service.ensure_assets(force=False)
        assets = self._create_support_images(summary)
        output = self._build_document(summary, assets)
        self.stdout.write(self.style.SUCCESS(f"Da tao file Word: {output}"))

    def _build_document(self, summary: dict[str, Any], assets: dict[str, Path]) -> Path:
        repo_root = Path(settings.BASE_DIR).parent
        output_path = repo_root / "aiservice02_lop.nhom_tenho.docx"
        document = Document()
        self._set_default_styles(document)
        self._add_cover_page(document)
        self._add_description_section(document, summary)
        self._add_dataset_section(document, summary, assets)
        self._add_model_section(document, summary, assets)
        self._add_graph_section(document, summary, assets)
        self._add_rag_section(document, summary, assets)
        document.save(output_path)
        return output_path

    def _set_default_styles(self, document: Document) -> None:
        normal = document.styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        document.sections[0].top_margin = Inches(0.8)
        document.sections[0].bottom_margin = Inches(0.8)
        document.sections[0].left_margin = Inches(0.8)
        document.sections[0].right_margin = Inches(0.8)

    def _add_cover_page(self, document: Document) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BÁO CÁO AISERVICE 02")
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(15, 23, 42)

        for text in [
            "",
            "Đề tài: Xây dựng AI Service cho hệ e-commerce",
            "Nội dung: Sinh dữ liệu, huấn luyện mô hình, KB_Graph, RAG và tích hợp chat/recommendation",
            "File nộp: aiservice02_lop.nhom_tenho.docx",
            "",
            "Sinh viên: ........................................................",
            "Lớp: .................................................................",
            "Nhóm: ..............................................................",
            "Giảng viên: ........................................................",
            "",
            "Ngày xuất tài liệu: 20/04/2026",
        ]:
            para = document.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run(text).font.size = Pt(14)

        document.add_page_break()

    def _add_description_section(self, document: Document, summary: dict[str, Any]) -> None:
        document.add_heading("1. Mô tả AISERVICE", level=1)
        paragraphs = [
            "AI Service được tích hợp trực tiếp vào dự án BTL theo mô hình microservice để phục vụ hai nhu cầu chính của hệ e-commerce: gợi ý sản phẩm theo hành vi người dùng và chat tư vấn mua sắm theo ngữ cảnh.",
            "Dữ liệu đầu vào của AI Service gồm dữ liệu hành vi người dùng user_id, product_id, action, timestamp. Từ dữ liệu này hệ thống sinh tập data_user500.csv, huấn luyện 3 mô hình tuần tự RNN, LSTM, BiLSTM, sau đó chọn model_best để dùng cho recommendation.",
            "Ngoài phần dự đoán hành vi, AI Service còn xây dựng Knowledge Base Graph để biểu diễn quan hệ User - Action - Product - Category - Brand. Từ graph này hệ thống tạo RAG cho chatbot nhằm trả lời theo ngữ cảnh catalog và tri thức hành vi.",
            "Trong hệ thống e-commerce, AI Service đã được nối với frontend để ghi nhận search, view, click, add_to_cart; đồng thời hiển thị danh sách gợi ý khi người dùng tìm kiếm hoặc thao tác với giỏ hàng và cung cấp giao diện chat riêng cho cửa hàng.",
        ]
        for text in paragraphs:
            document.add_paragraph(text)

        catalog = summary["catalog"]
        info_table = document.add_table(rows=5, cols=2)
        _table_style(info_table)
        data = [
            ("Số lượng sản phẩm catalog", str(catalog["products"])),
            ("Danh mục", ", ".join(catalog["categories"])),
            ("Thương hiệu", ", ".join(catalog["brands"])),
            ("Số bản ghi hành vi", str(summary["dataset"]["rows"])),
            ("Mô hình tốt nhất", summary["model_report"]["best_model"]),
        ]
        for row_idx, (label, value) in enumerate(data):
            _set_cell_text(info_table.rows[row_idx].cells[0], label, bold=True)
            _set_cell_text(info_table.rows[row_idx].cells[1], value)
        document.add_paragraph("")

    def _add_dataset_section(self, document: Document, summary: dict[str, Any], assets: dict[str, Path]) -> None:
        document.add_heading("2. Copy 20 dòng data", level=1)
        document.add_paragraph(
            "Tập dữ liệu được sinh tự động từ AI Service với 500 user, mỗi user có 8 hành vi, tổng cộng 4000 dòng. Ba hành vi dùng cho bài toán phân loại là view, click, add_to_cart."
        )
        document.add_picture(str(assets["action_distribution"]), width=Inches(6.2))
        cap = document.add_paragraph("Hình 1. Phân bố hành vi trong data_user500.csv")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rows = summary["dataset"]["sample_20_rows"]
        table = document.add_table(rows=len(rows) + 1, cols=4)
        _table_style(table)
        headers = ["user_id", "product_id", "action", "timestamp"]
        for col_idx, header in enumerate(headers):
            _set_cell_text(table.rows[0].cells[col_idx], header, bold=True)
            _shade_cell(table.rows[0].cells[col_idx], "DCE6F1")
        for row_idx, item in enumerate(rows, start=1):
            for col_idx, header in enumerate(headers):
                _set_cell_text(table.rows[row_idx].cells[col_idx], _safe_text(item[header]))

        document.add_paragraph(
            f"Đường dẫn file dữ liệu: {Path(settings.BASE_DIR) / summary['dataset']['path']}"
        )
        document.add_page_break()

    def _add_model_section(self, document: Document, summary: dict[str, Any], assets: dict[str, Path]) -> None:
        document.add_heading("3. Câu 2a - Huấn luyện RNN, LSTM, BiLSTM", level=1)
        document.add_paragraph(
            "Bài toán được thiết kế dưới dạng dự đoán và phân loại hành vi kế tiếp của user trên chuỗi tương tác. Hệ thống lấy lịch sử hành vi và thông tin category/brand của sản phẩm để dự đoán nhãn hành vi thuộc một trong ba lớp: view, click, add_to_cart."
        )
        document.add_paragraph(
            "Ba mô hình được triển khai gồm RNN, LSTM và BiLSTM. Sau khi huấn luyện, hệ thống đánh giá bằng accuracy, precision_macro, recall_macro và f1_macro. Mô hình có F1-macro cao nhất được chọn làm model_best."
        )

        document.add_picture(str(assets["model_comparison"]), width=Inches(6.2))
        cap = document.add_paragraph("Hình 2. So sánh F1-macro của ba mô hình")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        report = summary["model_report"]
        metrics_table = document.add_table(rows=len(report["models"]) + 1, cols=5)
        _table_style(metrics_table)
        headers = ["Mô hình", "Accuracy", "Precision", "Recall", "F1-macro"]
        for idx, header in enumerate(headers):
            _set_cell_text(metrics_table.rows[0].cells[idx], header, bold=True)
            _shade_cell(metrics_table.rows[0].cells[idx], "FCE4D6")
        for row_idx, model in enumerate(report["models"], start=1):
            metrics = model["metrics"]
            values = [
                model["name"],
                f"{metrics['accuracy']:.4f}",
                f"{metrics['precision_macro']:.4f}",
                f"{metrics['recall_macro']:.4f}",
                f"{metrics['f1_macro']:.4f}",
            ]
            for col_idx, value in enumerate(values):
                _set_cell_text(metrics_table.rows[row_idx].cells[col_idx], value, bold=(col_idx == 0 and value == report["best_model"]))

        document.add_paragraph(
            f"Kết luận chọn mô hình: {report['evaluation_note']}"
        )
        document.add_paragraph(
            "Nhận xét: RNN cho kết quả thấp hơn do chỉ nắm quan hệ gần. LSTM ghi nhớ chuỗi tốt hơn nên đạt kết quả cao nhất. BiLSTM cũng cho kết quả tốt nhưng vẫn thấp hơn LSTM trong tập test hiện tại."
        )

        document.add_heading("Code minh họa câu 2a", level=2)
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR) / "app" / "assignment.py",
                "class RNNSequenceModel",
                "class LSTMSequenceModel",
            ),
        )
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR) / "app" / "assignment.py",
                "def train_models",
                "def _evaluate_model",
            ),
        )
        document.add_paragraph(
            f"Ảnh kết quả chi tiết lưu tại: {Path(settings.BASE_DIR) / report['comparison_plot']}"
        )
        document.add_page_break()

    def _add_graph_section(self, document: Document, summary: dict[str, Any], assets: dict[str, Path]) -> None:
        document.add_heading("4. Câu 2b - Knowledge Base Graph với Neo4j", level=1)
        document.add_paragraph(
            "Knowledge Base Graph được xây dựng từ dữ liệu hành vi để biểu diễn các quan hệ giữa User, Action, Product, Category và Brand. Mỗi hành vi tạo cạnh từ user tới product, đồng thời liên kết thêm tới category và brand để tăng khả năng truy vấn ngữ cảnh."
        )
        document.add_paragraph(
            "Kết quả build graph có 547 node, 8341 edge và 4000 bản ghi nguồn. Hệ thống sinh đồng thời file JSON và file Cypher để có thể import vào Neo4j."
        )
        document.add_picture(str(assets["graph_overview"]), width=Inches(6.2))
        cap = document.add_paragraph("Hình 3. Minh họa cấu trúc KB_Graph")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        graph = json.loads((Path(settings.BASE_DIR) / summary["graph_report"]["path"]).read_text(encoding="utf-8"))
        edge_table = document.add_table(rows=21, cols=4)
        _table_style(edge_table)
        edge_headers = ["source", "relation", "target", "weight"]
        for idx, header in enumerate(edge_headers):
            _set_cell_text(edge_table.rows[0].cells[idx], header, bold=True)
            _shade_cell(edge_table.rows[0].cells[idx], "E2F0D9")
        for row_idx, edge in enumerate(graph["edges"][:20], start=1):
            values = [
                edge["source"],
                edge["relation"],
                edge["target"],
                _safe_text(edge["properties"].get("weight", 1)),
            ]
            for col_idx, value in enumerate(values):
                _set_cell_text(edge_table.rows[row_idx].cells[col_idx], value)

        document.add_heading("Code minh họa câu 2b", level=2)
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR) / "app" / "assignment.py",
                "def build_graph",
                "def load_graph",
            ),
        )
        document.add_paragraph(
            f"File graph JSON: {Path(settings.BASE_DIR) / summary['graph_report']['path']}"
        )
        document.add_paragraph(
            f"File Cypher import Neo4j: {Path(settings.BASE_DIR) / summary['graph_report']['cypher_path']}"
        )
        document.add_page_break()

    def _add_rag_section(self, document: Document, summary: dict[str, Any], assets: dict[str, Path]) -> None:
        document.add_heading("5. Câu 2c và 2d - RAG, Chat và tích hợp e-commerce", level=1)
        document.add_paragraph(
            "Phần RAG sử dụng hai nguồn ngữ cảnh chính: catalog sản phẩm và tri thức suy ra từ KB_Graph. Khi user gửi câu hỏi, hệ thống tách filter category, brand, giá, sau đó truy hồi các sản phẩm phù hợp và trích xuất thêm graph facts để đưa vào prompt của chatbot."
        )
        document.add_paragraph(
            "Trong bước tích hợp hệ e-commerce, frontend đã được nối thêm luồng theo dõi search, click, view, add_to_cart. RecommendationPanel hiển thị danh sách gợi ý khi user tìm kiếm hoặc đang thao tác ở giỏ hàng. AIChatBubble trở thành giao diện chat riêng của cửa hàng thay vì giao diện chat chung."
        )
        document.add_picture(str(assets["integration_overview"]), width=Inches(6.2))
        cap = document.add_paragraph("Hình 4. Sơ đồ tích hợp RAG và e-commerce")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

        rag_table = document.add_table(rows=5, cols=2)
        _table_style(rag_table)
        rows = [
            ("API chat", "/api/ai/chat"),
            ("API recommendations", "/api/ai/recommendations"),
            ("API assignment summary", "/api/ai/assignment-summary"),
            ("Frontend gợi ý search", "frontend/src/pages/HomePage.jsx"),
            ("Frontend chat", "frontend/src/components/AIChatBubble.jsx"),
        ]
        for idx, (label, value) in enumerate(rows):
            _set_cell_text(rag_table.rows[idx].cells[0], label, bold=True)
            _set_cell_text(rag_table.rows[idx].cells[1], value)

        document.add_heading("Code minh họa câu 2c", level=2)
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR) / "app" / "services.py",
                "class ChatRAGService",
                "    def _invoke_llm",
                include_end=False,
            ),
        )
        document.add_heading("Code minh họa câu 2d", level=2)
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR).parent / "frontend" / "src" / "components" / "RecommendationPanel.jsx",
                "export default function RecommendationPanel",
                "}",
                include_end=True,
            ),
        )
        self._add_code_block(
            document,
            self._extract_snippet(
                Path(settings.BASE_DIR).parent / "frontend" / "src" / "components" / "AIChatBubble.jsx",
                "export default function AIChatBubble",
                "}",
                include_end=True,
            ),
        )
        document.add_paragraph(
            "Kết luận: AI Service đã bao phủ đầy đủ chuỗi yêu cầu từ tạo dữ liệu, huấn luyện mô hình, biểu diễn tri thức bằng graph, truy hồi ngữ cảnh để chat, đến tích hợp trực tiếp trong trải nghiệm mua sắm của hệ thống."
        )

    def _extract_snippet(self, path: Path, start_marker: str, end_marker: str, include_end: bool = False) -> str:
        lines = path.read_text(encoding="utf-8").splitlines()
        start_idx = 0
        end_idx = len(lines)
        for idx, line in enumerate(lines):
            if start_marker in line:
                start_idx = idx
                break
        for idx in range(start_idx + 1, len(lines)):
            if end_marker in lines[idx]:
                end_idx = idx + (1 if include_end else 0)
                break
        snippet = "\n".join(lines[start_idx:end_idx]).strip()
        return snippet[:3500]

    def _add_code_block(self, document: Document, code: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        run = paragraph.add_run(code)
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    def _create_support_images(self, summary: dict[str, Any]) -> dict[str, Path]:
        reports_dir = Path(settings.BASE_DIR) / "artifacts" / "reports"
        images_dir = reports_dir / "doc_images"
        images_dir.mkdir(parents=True, exist_ok=True)

        action_distribution = images_dir / "action_distribution.png"
        model_comparison = images_dir / "model_comparison.png"
        graph_overview = images_dir / "graph_overview.png"
        integration_overview = images_dir / "integration_overview.png"

        self._draw_action_distribution(action_distribution, summary["dataset"]["action_distribution"])
        self._draw_model_comparison(model_comparison, summary["model_report"]["models"])
        graph = json.loads((Path(settings.BASE_DIR) / summary["graph_report"]["path"]).read_text(encoding="utf-8"))
        self._draw_graph_overview(graph_overview, graph)
        self._draw_integration_overview(integration_overview)

        return {
            "action_distribution": action_distribution,
            "model_comparison": model_comparison,
            "graph_overview": graph_overview,
            "integration_overview": integration_overview,
        }

    def _font(self, size: int, bold: bool = False):
        try:
            name = "arialbd.ttf" if bold else "arial.ttf"
            return ImageFont.truetype(name, size=size)
        except Exception:
            return ImageFont.load_default()

    def _canvas(self, width: int = 1600, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
        image = Image.new("RGB", (width, height), "#FFFFFF")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=28, outline="#CBD5E1", width=3, fill="#F8FAFC")
        return image, draw

    def _draw_action_distribution(self, path: Path, distribution: dict[str, int]) -> None:
        image, draw = self._canvas(1400, 820)
        title_font = self._font(48, True)
        text_font = self._font(28)
        draw.text((70, 55), "Phan bo hanh vi trong data_user500.csv", fill="#0F172A", font=title_font)
        total = sum(distribution.values()) or 1
        colors = {"view": "#2563EB", "click": "#F97316", "add_to_cart": "#16A34A"}
        bar_left = 120
        max_width = 900
        y = 180
        for action, value in distribution.items():
            draw.text((bar_left, y - 40), action, fill="#1E293B", font=self._font(30, True))
            width = int(max_width * value / total)
            draw.rounded_rectangle((bar_left, y, bar_left + width, y + 58), radius=18, fill=colors[action])
            draw.text((bar_left + width + 20, y + 10), f"{value} dong", fill="#334155", font=text_font)
            draw.text((1100, y + 10), f"{value / total * 100:.1f}%", fill="#334155", font=text_font)
            y += 150
        draw.text((120, 660), f"Tong so dong du lieu: {total}", fill="#475569", font=self._font(26, True))
        image.save(path)

    def _draw_model_comparison(self, path: Path, models: list[dict[str, Any]]) -> None:
        image, draw = self._canvas(1500, 900)
        draw.text((70, 55), "So sanh 3 mo hinh RNN - LSTM - BiLSTM", fill="#0F172A", font=self._font(48, True))
        draw.text((70, 118), "Gia tri cot the hien F1-macro tren tap test", fill="#475569", font=self._font(24))
        base_y = 720
        draw.line((130, base_y, 1320, base_y), fill="#94A3B8", width=4)
        colors = ["#2563EB", "#0F766E", "#7C3AED"]
        x = 220
        for idx, model in enumerate(models):
            f1 = float(model["metrics"]["f1_macro"])
            bar_height = int(430 * f1)
            draw.rounded_rectangle((x, base_y - bar_height, x + 180, base_y), radius=22, fill=colors[idx])
            draw.text((x + 42, base_y + 20), model["name"], fill="#1E293B", font=self._font(30, True))
            draw.text((x + 45, base_y - bar_height - 50), f"{f1:.3f}", fill="#1E293B", font=self._font(28, True))
            x += 320
        image.save(path)

    def _draw_graph_overview(self, path: Path, graph: dict[str, Any]) -> None:
        image, draw = self._canvas(1500, 900)
        draw.text((70, 55), "Mo hinh KB_Graph", fill="#0F172A", font=self._font(48, True))
        draw.text((70, 118), "Quan he giua User - Action - Product - Category - Brand", fill="#475569", font=self._font(24))
        positions = {
            "User": (220, 430),
            "Action": (560, 220),
            "Product": (890, 430),
            "Category": (1220, 240),
            "Brand": (1220, 620),
        }
        counts = Counter(node["label"] for node in graph["nodes"])
        colors = {
            "User": "#2563EB",
            "Action": "#F97316",
            "Product": "#16A34A",
            "Category": "#7C3AED",
            "Brand": "#DC2626",
        }
        edges = [("User", "Action", "PERFORMED"), ("User", "Product", "view/click/add_to_cart"), ("Product", "Category", "IN_CATEGORY"), ("Product", "Brand", "OF_BRAND"), ("User", "Category", "INTERESTED_IN")]
        for source, target, label in edges:
            sx, sy = positions[source]
            tx, ty = positions[target]
            draw.line((sx, sy, tx, ty), fill="#94A3B8", width=5)
            mx, my = (sx + tx) / 2, (sy + ty) / 2
            draw.rounded_rectangle((mx - 110, my - 22, mx + 110, my + 22), radius=14, fill="#FFFFFF", outline="#CBD5E1")
            draw.text((mx - 95, my - 10), label, fill="#334155", font=self._font(18, True))
        for label, (x, y) in positions.items():
            r = 90 if label == "Product" else 78
            draw.ellipse((x - r, y - r, x + r, y + r), fill=colors[label], outline="#0F172A", width=3)
            draw.text((x - 45, y - 18), label, fill="#FFFFFF", font=self._font(26, True))
            draw.text((x - 60, y + 20), f"{counts.get(label, 0)} node", fill="#FFFFFF", font=self._font(18))
        stats = graph["stats"]
        draw.rounded_rectangle((70, 720, 660, 820), radius=20, fill="#E2E8F0")
        draw.text((100, 750), f"Thong ke: {stats['node_count']} node | {stats['edge_count']} edge | {stats['rows']} dong du lieu", fill="#1E293B", font=self._font(24, True))
        image.save(path)

    def _draw_integration_overview(self, path: Path) -> None:
        image, draw = self._canvas(1500, 950)
        draw.text((70, 55), "Tich hop RAG va he e-commerce", fill="#0F172A", font=self._font(48, True))
        draw.text((70, 118), "Search, click, view, add_to_cart duoc dua vao AI Service de goi y va chat", fill="#475569", font=self._font(24))
        boxes = [
            ((120, 280, 420, 430), "#DBEAFE", "Frontend", "HomePage, CartPage,\nAIChatBubble"),
            ((560, 280, 920, 430), "#FDE68A", "API Gateway", "Forward /api/ai/chat,\n/api/ai/recommendations"),
            ((1060, 180, 1420, 350), "#DCFCE7", "AI Service", "Track event,\nmodel_best,\nRAG chat"),
            ((1060, 470, 1420, 640), "#E9D5FF", "KB_Graph", "User-Product-\nCategory-Brand"),
            ((560, 620, 920, 790), "#FECACA", "Product Catalog", "Mobile, Desktop,\nClothes service"),
        ]
        for box, color, title, subtitle in boxes:
            draw.rounded_rectangle(box, radius=28, fill=color, outline="#94A3B8", width=3)
            draw.text((box[0] + 26, box[1] + 24), title, fill="#0F172A", font=self._font(32, True))
            draw.multiline_text((box[0] + 26, box[1] + 76), subtitle, fill="#334155", font=self._font(24), spacing=8)
        arrows = [
            ((420, 355), (560, 355), "request"),
            ((920, 355), (1060, 265), "chat/reco"),
            ((920, 355), (1060, 555), "graph sync"),
            ((920, 705), (1060, 555), "catalog context"),
        ]
        for start, end, label in arrows:
            draw.line((*start, *end), fill="#475569", width=5)
            angle = math.atan2(end[1] - start[1], end[0] - start[0])
            arrow_size = 16
            x3 = end[0] - arrow_size * math.cos(angle - math.pi / 6)
            y3 = end[1] - arrow_size * math.sin(angle - math.pi / 6)
            x4 = end[0] - arrow_size * math.cos(angle + math.pi / 6)
            y4 = end[1] - arrow_size * math.sin(angle + math.pi / 6)
            draw.polygon([end, (x3, y3), (x4, y4)], fill="#475569")
            mx, my = (start[0] + end[0]) / 2, (start[1] + end[1]) / 2
            draw.text((mx - 40, my - 30), label, fill="#1E293B", font=self._font(18, True))
        image.save(path)
